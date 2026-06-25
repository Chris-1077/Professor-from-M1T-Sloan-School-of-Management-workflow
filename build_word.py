"""
Parse the full PDF text and generate a well-formatted Word document.
"""
import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# STEP 1: Parse the full text
# ============================================================
with open('D:/ENLIVE/temp_pdf_full.txt', 'r', encoding='utf-8') as f:
    full_text = f.read()

# Split by pages
pages = re.split(r'=== PAGE \d+ ===\n', full_text)
pages = [p.strip() for p in pages if p.strip()]
print(f"Total pages: {len(pages)}")

# ============================================================
# STEP 2: Identify content pages (skip question-only pages)
# ============================================================
# A page is a "question page" if it primarily contains 例题, 答案, or 解析 without substantive new content
question_patterns = [
    r'^【(单选题|多选题|判断题|填空题|名词解释|简答题|论述题)】',
    r'^【答案】',
    r'^【解析】',
    r'^【参考答案】',
]

def is_question_page(page_text):
    """Check if page is primarily a question/answer page (not new content)."""
    lines = [l.strip() for l in page_text.split('\n') if l.strip()]
    if not lines:
        return True

    # Count question/answer lines vs content lines
    qa_count = 0
    content_count = 0
    for line in lines:
        if any(re.match(p, line) for p in question_patterns):
            qa_count += 1
        elif re.match(r'^(第[一二三四五六七八九十]+章|考点\d|本章|课程知|题型|自学|应试|谢谢)', line):
            content_count += 1
        elif len(line) > 20 and ('含义' in line or '是指' in line or '包括' in line):
            content_count += 1

    # If mostly QA, skip
    if qa_count > content_count and qa_count >= 2:
        return True
    # If page only has one question and answer
    if qa_count >= 1 and content_count == 0:
        return True
    return False

# Also skip pages that are just "谢谢" or section dividers
def is_skip_page(page_text):
    lines = [l.strip() for l in page_text.split('\n') if l.strip()]
    if not lines:
        return True
    if len(lines) <= 2 and any(l in ['谢谢', '考试分析', '考试分析 '] for l in lines):
        return True
    return False

# ============================================================
# STEP 3: Extract structured content
# ============================================================
chapters = {}
current_chapter = None
current_kaodian = None
current_section = None

chapter_names = {
    '第一章': '购买、拥有和存在：消费者行为学概述',
    '第二章': '消费者与社会福祉',
    '第三章': '知觉',
    '第四章': '学习和记忆',
    '第五章': '动机和情感',
    '第六章': '自我：心智、性别和身体',
    '第七章': '个性、生活方式和价值观',
    '第八章': '态度与劝说沟通',
    '第九章': '制定决策',
    '第十章': '购买、使用与处置',
    '第十一章': '群体和社交媒体',
    '第十二章': '收入和社会阶层',
    '第十三章': '亚文化',
    '第十四章': '文化',
}

key_chapters = {'第三章', '第四章', '第五章', '第八章', '第十章'}

# Process content pages
content_pages = []
for i, page in enumerate(pages):
    if not is_skip_page(page) and not is_question_page(page):
        content_pages.append((i+1, page))

print(f"Content pages (after filtering): {len(content_pages)}")

# ============================================================
# STEP 4: Build chapter-based knowledge structure
# ============================================================
# We'll organize by chapter, then by kaodian (exam point)

chapter_pattern = re.compile(r'(第[一二三四五六七八九十]+章)\s*(.*)')
kaodian_pattern = re.compile(r'(考点\d+)\s*(.+)')

# Extract all kaodian headers and their content
kaodian_content = {}
current_kd = None
current_kd_lines = []

chapter_meta = {}  # chapter_num -> {'kaodians': [kd1, kd2, ...]}

for page_num, text in content_pages:
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Check for chapter
        ch_match = chapter_pattern.match(line)
        if ch_match:
            ch_num = ch_match.group(1)
            ch_title = ch_match.group(2) if ch_match.group(2) else chapter_names.get(ch_num, '')
            if ch_num not in chapters:
                chapters[ch_num] = {'title': ch_title, 'kaodians': [], 'key': ch_num in key_chapters}
                chapter_meta[ch_num] = {'kaodians': []}
            continue

        # Check for kaodian
        kd_match = kaodian_pattern.match(line)
        if kd_match:
            # Save previous kaodian
            if current_kd and current_kd_lines:
                kaodian_content[current_kd] = '\n'.join(current_kd_lines)

            kd_num = kd_match.group(1)
            kd_title = kd_match.group(2)
            current_kd = kd_num
            current_kd_lines = [kd_title]

            # Determine which chapter this belongs to
            for ch_num in reversed(list(chapters.keys())):
                if current_kd not in chapter_meta.get(ch_num, {}).get('kaodians', []):
                    chapter_meta.setdefault(ch_num, {'kaodians': []})
                    chapter_meta[ch_num]['kaodians'].append(current_kd)
                    chapters[ch_num]['kaodians'].append({'id': current_kd, 'title': kd_title})
                break
            continue

        # Add line to current kaodian
        if current_kd:
            current_kd_lines.append(line)

# Save last kaodian
if current_kd and current_kd_lines:
    kaodian_content[current_kd] = '\n'.join(current_kd_lines)

# ============================================================
# STEP 5: Manually curate the knowledge content
# ============================================================
# Since automated parsing of PDF text layout is unreliable for complex formatting,
# we'll build the knowledge base directly from the raw text, organized by chapter.

KNOWLEDGE = {
    '第一章': {
        'title': '购买、拥有和存在：消费者行为学概述',
        'key': False,
        'points': [
            {
                'title': '一、消费者行为学的基本概念',
                'items': [
                    '消费者行为学：研究个体或群体为满足需要与欲望而挑选、购买、使用或处置产品、服务、观念或经验所涉及的过程。',
                    '消费者行为是一个持续的过程，包括购买前、购买中和购买后影响消费者的所有问题。其中重要的部分是交易——两个或两个以上组织或个人付出和取得某种有价值的东西的过程。',
                    '消费者：在消费过程的三个阶段中，产生需要或欲望、实施购买并处置产品的人。注意：购买者和使用者可能不是同一个人。',
                ]
            },
            {
                'title': '二、市场中的人们',
                'items': [
                    '人口统计特征：描述一个群体特征的变量，包括年龄、性别、收入或职业等。',
                    '网上消费社区：互联网的发展已造就了成千上万的网上消费社区，成员们在那里分享各种各样的内容。',
                    '市场细分战略：企业将其产品、服务和想法只提供给特定的消费者群，而非所有人。',
                    '品牌：由广告、包装、品牌化策略，以及其他营销要素共同塑造，并且是规划好的清晰形象或"个性"。',
                    '80/20法则：20%的用户贡献了总销量的80%。营销者用来识别最忠实的顾客或频繁使用者。',
                ]
            },
            {
                'title': '三、消费者对营销策略的影响',
                'items': [
                    '区分消费者的变量：人口统计变量（年龄、性别、家庭结构、社会阶层与收入、民族和种族）；地理位置变量；生活方式变量；行为细分变量。',
                    '关系营销：企业成功的关键因素之一——在品牌与消费者间建立起持续终身的关系。营销者会与消费者定期沟通，说服他们与企业保持长期关系。在经济形势不好的情况下，关系营销更加重要。',
                    '数据库营销：密切跟踪消费者的购物习惯，并根据这些信息设计适合人们欲望与需要的产品和广告。',
                    '大数据：海量数据库的集合与分析。',
                ]
            },
            {
                'title': '四、营销对消费者的影响',
                'items': [
                    '通俗文化：包括音乐、电影、运动、图书、名人和其他大众市场消费的娱乐形式，对营销者而言既是产品也是灵感。',
                    '角色理论：大部分消费者行为都如同在演一场戏，每个消费者都有特定的台词、道具和服装。因为人们要扮演许多不同的角色，有时需要根据当下正在演出的"戏"来调整消费决策。',
                    '需要（need）：一个人为了生存或实现某一目标所必须有的东西。欲望（want）：某种需要的具体呈现，由个人或文化因素决定。',
                    '人与产品的关系类型：①自我概念依附——产品有助于确立使用者的身份；②怀旧依附——产品成为与过去的自我的一种联结；③互相依附——产品成为使用者日常生活的一部分；④爱——产品成为引发温暖、激情或其他强烈情绪的情感纽带。',
                ]
            },
            {
                'title': '五、数字化居民与社会化媒体',
                'items': [
                    '商务革命：B2C（企业对消费者）、C2C（消费者对消费者）。',
                    '数字居民：生活在一个高度网络化且从不掉线的世界，数字技术在其成长过程中一直存在着。',
                    '横向革命：在某种程度上以社交媒体的普及为标志。被技术性能和流动性武装起来的人、社区和组织组成相互联系和相互依赖的网络。',
                    '社交媒体：在网络中在线交流、传送、协作和建立关系的手段。',
                ]
            },
            {
                'title': '六、消费者研究的两种观念',
                'items': [
                    '实证主义（positivism）：人类的理性至高无上，认为存在能被科学发现的唯一的客观真理。重视客体的功能、推崇科技，将世界视为一个理性而有序的场所。',
                    '解释主义（interpretivism）：认为科学技术被过度强调，有序、理性的观点否认了社会与文化的复杂性。强调象征性主观经验的重要性，强调意义存在于个人意识中。产生的知识受时间限制，依赖背景。',
                ]
            },
        ]
    },

    '第二章': {
        'title': '消费者与社会福祉',
        'key': False,
        'points': [
            {
                'title': '一、商业伦理与消费者权益',
                'items': [
                    '商业伦理：引导市场行为的规则——一种文化中的大多数人判断好坏、对错的标准。通用标准包括诚实、可信赖、公平、尊重、正义感、正直、关心他人、责任感和忠诚。',
                    '营销是否创造人为需要？——营销的基本目标是对已存在的需要产生觉察，而非创造需要。',
                    '营销是否必需？——产品只是满足了人们既有的需要，营销行为仅仅是帮助人们与想要获取的东西保持沟通。广告帮助人们减少搜寻时间，是消费者愿意付费的服务。',
                    '营销者并未承诺奇迹——从新产品高达40%-80%的失败率可以得到印证，营销者对人们的了解还远未达到能操纵他们的程度。',
                    '物质主义：人们对世俗财物的重视程度。"你所拥有的就是你吗？"',
                ]
            },
            {
                'title': '二、消费者权益和产品满意度',
                'items': [
                    '消费者不满时的三种行动选择：①语言反应——向零售商提出赔偿诉求；②私下反应——将不满告诉朋友并抵制购买该产品，或换到其他商店购买；③求助第三方——提起法律诉讼、向消费者保护组织投诉或写信给媒体。',
                    '消费者基本权利（《消费者权利宣言》）：安全权、知情权、获赔权、选择权。',
                    '变革式消费研究：从事消费者研究的人自己组织起来，研究并纠正在市场上观察到的一些严重的社会问题。',
                    '企业社会责任（CSR）：鼓励组织对包括消费者、雇员和环境在内的不同利益相关者产生正面影响的过程。',
                ]
            },
            {
                'title': '三、与消费者行为相关的主要政策问题',
                'items': [
                    '数据隐私和身份盗用：身份盗用——未经允许而窃取和使用你的个人信息；僵尸网络攻击；位置隐私——手机随时发送其大体位置。',
                    '市场进入限制：很多人不能自由地寻找和购买商品或服务。原因包括残障、食品荒漠、媒体素养。',
                    '三重底线导向商业战略：①财务底线——向股东提供利润；②社会底线——将利益回馈所在的社区；③环境底线——最小化对环境的破坏，尽量改善自然条件。',
                    '绿色营销：对环境友好型产品的开发和推广，并且制造商把产品卖给消费者时会特意强化这一特征。',
                ]
            },
            {
                'title': '四、消费者行为的黑暗面',
                'items': [
                    '消费者恐怖主义：金融崩溃、电子和供应网络中断；生物恐怖主义和网络恐怖主义。',
                    '成瘾消费：对于产品或者服务产生的生理或者心理依赖。社交媒体成瘾是其中一种。',
                    '强迫消费：一种因为想缓解焦虑、沮丧或厌倦而重复和经常过度地消费。特征：①行动没有选择性；②来自行为的满足是暂时的；③过后常常会有强烈的后悔和愧疚感。',
                    '被消费的消费者：为了在市场上获得商业利益，而使得一些人被使用或被利用，不论其是否愿意。',
                    '非法产品获取和使用：消费者盗窃和欺诈、伪造、反消费。',
                ]
            },
        ]
    },

    '第三章': {
        'title': '知觉',
        'key': True,
        'points': [
            {
                'title': '一、感觉系统',
                'items': [
                    '感觉：我们的感受器（眼、耳、鼻、口、指和皮肤）对光线、色彩、声音、气味和质地等基本刺激的直接反应。',
                    '知觉：对这些感觉进行选择、组织和解释的过程。',
                    '享乐性消费：消费者与产品相互作用的多重感官感觉、幻想和情感方面。',
                    '感官营销：特别注重感官影响产品体验，包括视觉、气味、声音、触觉、味觉五个方面。',
                ]
            },
            {
                'title': '二、知觉的三个阶段（暴露 → 注意 → 解释）',
                'items': [],
            },
            {
                'title': '第一阶段：暴露',
                'items': [
                    '暴露：当一个刺激进入个人感觉器官接收范围之内时，暴露就发生了。',
                    '感觉阈限：一个临界点，只有超过这个临界点的刺激，才足以对一个人的感觉产生被意识到的影响。',
                    '绝对阈限：特定感觉渠道所能觉察到的最小刺激量。',
                    '差别阈限：感觉系统觉察两种刺激之间的差别或者变化的能力。',
                    '最小可觉察差别（JND）：能够觉察到的两种刺激之间的最小差别。',
                    '韦伯定律：引起注意所需要的刺激变化量与初始刺激强度有关。初始刺激越强，引起注意所需要的刺激变化量越大。',
                    '阈下知觉：刺激在消费者的感知水平之下。',
                ]
            },
            {
                'title': '第二阶段：注意',
                'items': [
                    '注意：对特定刺激进行信息加工的投入程度。',
                    '广告混战：在刺激信息日益喧嚣的环境中，过多营销信息争夺注意力。',
                    '认知选择（知觉选择）：消费者有选择地注意信息，只注意所接触刺激的一小部分，实践"心理经济"，以免被信息淹没。',
                    '个人选择因素：①知觉警惕——消费者更可能意识到与当前需要有关的刺激物；②知觉防御——人们倾向于看见想要看的东西，而看不见不想看的东西，若刺激构成威胁，可能不理会或改变其意义；③适应——消费者对某个刺激持续注意的程度，当过于熟悉而不再注意时，适应就发生了。',
                    '导致适应的因素：①强度——较低强度刺激容易出现适应；②甄别——简单刺激易于适应；③暴露——频繁接触的刺激更容易适应；④关联性——无关或不重要的刺激会引起适应。',
                    '刺激物选择因素：与周围其他刺激物形成对比的刺激更可能引起注意。对比的产生方式：大小、颜色、位置、新颖性。',
                ]
            },
            {
                'title': '第三阶段：解释',
                'items': [
                    '解释：人们赋予感觉刺激物的意义。消费者赋予某一刺激物的意义取决于心理图式——人们赋予这一刺激物的信念集合。',
                    '格式塔心理学——刺激物组织方式的三原则：①闭合原则——人们倾向于把不完整的图形感知为完整的图形；②相似原则——消费者倾向于将已有相似物理特性的物品归类在一起；③图形-背景原则——刺激物的一个部分居于图形主导地位，而其他部分是相对次要的背景。',
                    '符号学：研究标志和象征之间的联系，以及它们在赋予意义中的作用。三个基本要素：①目标客体——广告语所指向的产品；②标志（形象）——能够传达意欲含义的感官形象；③诠释（含义）——消费者引申的含义。',
                    '知觉定位：人们对产品的评价一般是根据产品本身的意义，而不是产品能够做什么。这种消费者感知的意义构成了产品的市场定位。营销者可以利用多个维度来建立品牌定位：生活方式、价格领先、属性、产品类别、竞争者、场合、使用者、质量。',
                ]
            },
        ]
    },

    '第四章': {
        'title': '学习和记忆',
        'key': True,
        'points': [
            {
                'title': '一、学习的基本概念',
                'items': [
                    '学习：由经验引起的相对比较长久的行为改变。学习者不一定直接获得经验，也可以通过观察那些对他人产生影响的事件获得经验。',
                    '无意识学习：有时甚至不做任何尝试也在学习，随意的、无意识的知识获取过程。',
                    '学习理论主要包括行为理论和认知理论两大类。',
                ]
            },
            {
                'title': '二、行为学习理论——经典条件反射',
                'items': [
                    '经典条件反射：将一种能够诱发某种反应的刺激与另一种原本不能单独诱发这种反应的刺激相配对，随着时间的推移，第二种刺激会引起类似的反应。由苏联生物学家伊凡·巴甫洛夫在狗身上首次发现。',
                    '重复：在条件刺激和非条件刺激多次相互配对之后，条件作用效果就更有可能发生。重复增强了刺激和反应之间的联结，并防止这种联结在记忆中淡化。',
                    '消退：如果条件刺激和非条件刺激只是偶尔配对，条件反射就不会发生或不会持久发生。以前的条件作用逐渐减弱并最终消失。',
                    '刺激泛化：与条件刺激相似的刺激往往会引起类似的条件反应。',
                    '刺激甄别：当受到一种类似于条件刺激的刺激时，非条件刺激的行为并不发生。当刺激甄别启动时，反应会减弱并很快消失。',
                    '经典条件反射的营销应用——基于刺激泛化的四种策略：①家族品牌——许多产品因公司良好的品牌声誉而盈利；②产品线延伸——利用原有品牌再推出相关产品；③许可——公司把知名品牌名称"租赁"给别人使用；④相似包装——通过独特的包装设计产生对某一个特定品牌的强烈联想。',
                ]
            },
            {
                'title': '三、行为学习理论——工具性（操作性）条件反射',
                'items': [
                    '工具性条件反射：个体学会那些能产生积极结果并避免负面结果的行为。由心理学家B.F.斯金纳研究与这一学习过程联系最密切。',
                    '发生方式：①正强化——环境通过给予奖励的方式提供正强化，反应得到加强并使个体学会适当的行为；②负强化——为了避免不愉快而做某些事情；③惩罚——不愉快事件发生后的反应，通过难堪的事件学会不要重复这些行为；④消退——当不再获得积极的结果时。',
                    '营销应用：频繁营销——用奖品来激励老顾客，奖品随着购物数量的增加而增加。游戏化策略——通过将游戏元素添加到原本比较枯燥的任务中，将日常行为转化为一种体验。',
                ]
            },
            {
                'title': '四、认知学习理论',
                'items': [
                    '认知理论：将消费者看作复杂问题的解决者，他们通过观察他人来学习抽象的规则和概念。',
                    '观察学习：通过观察他人的行为和行为的结果来学习。',
                    '社会化：消费者学习成为社会成员的过程。对年幼儿童而言，主要来源是家庭和媒体。',
                ]
            },
            {
                'title': '五、记忆',
                'items': [
                    '记忆过程：感觉记忆 → 短时记忆（工作记忆）→ 长时记忆。',
                    '感觉记忆：对感觉信息的短暂储存，容量大但持续时间极短（不到1秒）。',
                    '短时记忆（工作记忆）：在有限时间内（约20秒）储存有限信息（约7±2个组块）。',
                    '长时记忆：信息可以被永久储存，容量几乎是无限的。',
                    '编码：信息从短时记忆进入长时记忆的过程。',
                    '提取：从长时记忆中调取信息的过程。',
                    '品牌联想：通过独特的包装设计产生对某一个特定品牌的强烈联想。',
                ]
            },
            {
                'title': '六、怀旧的营销力量',
                'items': [
                    '怀旧营销：利用消费者对过去的美好回忆来推广产品。怀旧能唤起温暖、舒适和安全感，增强品牌忠诚度。',
                    '怀旧依附：产品成为与过去的自我的一种联结。',
                ]
            },
        ]
    },

    '第五章': {
        'title': '动机和情感',
        'key': True,
        'points': [
            {
                'title': '一、动机的基本概念',
                'items': [
                    '动机：引导和维持人们行为的内部动力。动机来源于需要。',
                    '需要 vs 需求：需要是个体缺乏时的主观状态（如对食物的生理需要）；需求是需要与购买力和购买意愿相结合时产生的。',
                    '驱动理论：生理需要产生紧张状态，驱动人们采取行动来降低紧张。',
                    '期望理论：行为不仅由需要驱动，还由对结果的期望驱动。',
                ]
            },
            {
                'title': '二、马斯洛需求层次理论',
                'items': [
                    '五个层次（由低到高）：①生理需求（食物、水、睡眠）；②安全需求（安全、稳定、免受伤害）；③社交需求（归属感、爱、友谊）；④尊重需求（自尊、认可、地位）；⑤自我实现需求（发挥潜能、实现理想）。',
                    '核心观点：只有低层次需求得到基本满足后，更高层次的需求才会成为行为的主要驱动力。',
                ]
            },
            {
                'title': '三、消费者目标与隐性需求',
                'items': [
                    '消费者目标：短期/长期目标、具体/抽象目标。目标对购买决策有导向作用。',
                    '隐性需求：消费者自身可能难以清晰表达的需求，常源于潜意识或社会文化因素。企业可通过市场调研和大数据发现隐性需求。',
                    '技术发展改变需求：智能手机带动移动支付、短视频娱乐等新需求；消费动机从物质满足向社交、自我实现转变。',
                ]
            },
            {
                'title': '四、动机冲突',
                'items': [
                    '双趋冲突：两个目标都具有吸引力，但只能选择一个。',
                    '双避冲突：两个都不想面对的选择，但必须选一个。',
                    '趋避冲突：同一目标兼具吸引和排斥两方面的特征。',
                ]
            },
            {
                'title': '五、情感与消费',
                'items': [
                    '情绪对消费的影响：积极情绪促进购买，对价格敏感度降低；消极情绪可能导致推迟或取消购买。',
                    '情绪性消费：通过消费行为调节自身情绪。负面情绪时购物获得愉悦感；积极情绪时消费庆祝。',
                    '情感营销：品牌通过唤起消费者特定情感（如温暖、怀旧、兴奋）来建立情感联结。',
                ]
            },
        ]
    },

    '第六章': {
        'title': '自我：心智、性别和身体',
        'key': False,
        'points': [
            {
                'title': '一、自我概念',
                'items': [
                    '自我概念：个体对自身的认知与看法，包括实际自我、理想自我、社会自我等维度。',
                    '自我意象：个体期望呈现给外界的形象。',
                    '自尊水平：高自尊者更自信冒险，愿意尝试新品牌新产品；低自尊者更保守，依赖他人意见。',
                    '消费与自我认同：消费者通过消费选择来认识自我和表达自我。',
                ]
            },
            {
                'title': '二、身份与消费',
                'items': [
                    '身份影响消费：不同身份（职业、性别、年龄、社会阶层）具有不同消费偏好。如职场白领偏好简约正式风格；学生群体注重性价比。',
                    '性别角色与消费：男性与女性在消费行为上的差异，以及性别角色变化对营销的影响。',
                    '身体形象：消费者对自身身体形象的认知影响其购买行为（如健身产品、美容产品）。',
                    '理想自我营销：营销者常利用消费者对理想自我的向往来推广产品。',
                ]
            },
        ]
    },

    '第七章': {
        'title': '个性、生活方式和价值观',
        'key': False,
        'points': [
            {
                'title': '一、个性特质',
                'items': [
                    '弗洛伊德人格结构理论：本我（与即时满足有关，遵循快乐原则）、超我（代表道德规范，遵循理想原则）、自我（在两者间平衡，遵循现实原则）。',
                    '大五人格模型：①开放性（对新经验的开放程度）；②责任感（自律和尽责程度）；③外向性（社交活跃程度）；④亲和性（对他人的尊重和关心程度）；⑤神经质（情绪稳定程度）。',
                    '品牌个性：人们赋予品牌的一系列拟人化特质。品牌也可以被消费者感知为具有"个性"。',
                ]
            },
            {
                'title': '二、生活方式与价值观',
                'items': [
                    '生活方式：个体如何分配时间和金钱的生活模式，包括活动（Activities）、兴趣（Interests）和意见（Opinions），简称AIO。',
                    'VALS框架：基于价值观和生活方式将消费者分类的系统（Values and Lifestyles）。',
                    '价值观：持久的信念，影响消费者的长期行为模式。',
                    '文化适应（enculturation）：学习自己文化所认可的信念与行为的过程。',
                    '文化融合（acculturation）：学习另一种文化所认可的信念与行为的过程。',
                ]
            },
        ]
    },

    '第八章': {
        'title': '态度与劝说沟通',
        'key': True,
        'points': [
            {
                'title': '一、态度的定义与ABC模型',
                'items': [
                    '态度：对某一对象持有的持续评价性反应，包括认知（信念）、情感（感觉）、行为（行动）三种成分。',
                    'ABC态度模型：认知（Cognition / Belief）→ 情感（Affect / Feeling）→ 行为（Behavior / Action），三者相互影响。',
                    '态度一致性：消费者倾向于保持认知、情感和行为之间的一致性。',
                ]
            },
            {
                'title': '二、态度形成与改变',
                'items': [
                    '多属性态度模型：消费者根据产品属性及其重要性来形成态度。每个属性有其权重和评价分数，加权求和得到总体态度。',
                    '态度改变策略：①改变信念——改变消费者对产品属性的认知；②改变重要性权重——让消费者更重视某属性；③增加新信念——为产品增加新的属性维度。',
                ]
            },
            {
                'title': '三、精细加工可能性模型（ELM）',
                'items': [
                    '中心路径：基于理性思考，消费者认真分析信息内容，适用于高卷入度产品。通过中心路径形成的态度更持久、更抗改变。',
                    '边缘路径：基于情感和表面线索（如代言人吸引力、音乐、画面美），适用于低卷入度产品。通过边缘路径形成的态度较短暂。',
                    '影响路径选择的因素：消费者的信息处理能力、动机和卷入度。',
                ]
            },
            {
                'title': '四、劝说沟通',
                'items': [
                    '劝说沟通要素：①信源（可信度、吸引力）；②信息（内容质量、情感诉求方式）；③渠道（媒介类型）；④受众（目标群体的特征）。',
                    '信源可信度：专家性和可信赖性。高可信度信源更有效。',
                    '信源吸引力：外表吸引力、相似性、喜爱度。',
                    '恐惧诉求：适度的恐惧能促进态度改变，但过强的恐惧可能产生防御性回避。',
                    '双面说服：同时呈现正反两面论点，对教育程度高的受众更有效。',
                ]
            },
            {
                'title': '五、认知失调理论',
                'items': [
                    '认知失调：当行为与态度不一致时产生的不适感，消费者会调整态度或行为来减少失调。',
                    '购后认知失调：购买后怀疑自己的决策是否正确。营销者需要通过售后服务、正面评价等方式帮助消费者减少失调。',
                ]
            },
        ]
    },

    '第九章': {
        'title': '制定决策',
        'key': False,
        'points': [
            {
                'title': '一、消费者决策过程',
                'items': [
                    '五阶段模型：①问题识别 → ②信息搜索 → ③备选方案评估 → ④购买决策 → ⑤购后行为。',
                    '问题识别：消费者意识到理想状态与实际状态之间存在差距。',
                    '信息搜索来源：内部搜索（已有记忆和经验）和外部搜索（个人来源、商业来源、公共来源、经验来源）。',
                    '评价标准：消费者用来比较产品的维度和属性。不同消费者的评价标准可能不同。',
                ]
            },
            {
                'title': '二、决策类型',
                'items': [
                    '扩展型决策：高卷入度、高风险，消费者花大量时间搜索和比较。如购车、购房。',
                    '有限型决策：中等卷入度，消费者使用简单的决策规则。如更换品牌尝试。',
                    '习惯型决策：低卷入度，消费者几乎不进行信息搜索，自动做出选择。如日常食品购买。',
                ]
            },
            {
                'title': '三、启发式决策与偏差',
                'items': [
                    '启发式决策：心理捷径，帮助消费者快速做出决策。',
                    '品牌效应：知名品牌被认为质量更好。',
                    '价格-质量推论：高价格常被当作高质量的信号。',
                    '国家来源效应：产品的原产国影响质量感知（如德国汽车、瑞士手表）。',
                    '框架效应：问题的呈现方式影响消费者的决策。',
                    '锚定效应：第一印象或初始信息对后续判断产生不成比例的影响。',
                ]
            },
        ]
    },

    '第十章': {
        'title': '购买、使用与处置',
        'key': True,
        'points': [
            {
                'title': '一、情境因素',
                'items': [
                    '情境因素包括：①物理环境（装潢、音乐、灯光、气味）；②社交环境（他人在场及其影响）；③时间因素（可支配时间、时机）；④任务定义（购买目的——自用还是送礼）；⑤先前状态（心情、经济状况）。',
                ]
            },
            {
                'title': '二、购物动机',
                'items': [
                    '实用型动机：以功能性目的为主，注重效率和经济性。',
                    '享乐型动机：以愉快体验为目的，包括：①社交体验——购物作为社交活动；②分享共同爱好——在兴趣社群中消费；③归属感——通过消费融入群体；④地位——通过消费展示身份；⑤狩猎的兴奋感——寻找折扣和好货的乐趣。',
                    '主题化技巧：富有创造性的卖家使用的四种主题——风景主题、市场主题、网络空间主题和心境主题。',
                ]
            },
            {
                'title': '三、店铺选择与购物体验',
                'items': [
                    '店铺选择因素：位置、商品种类、价格、服务质量、店铺氛围。',
                    '购物体验管理：通过场景设计（与品牌定位契合）、互动参与（提升投入感）和情感共鸣（核心）来提升体验。',
                    '体验营销的风险：标准化难、成本高、期望管理难。',
                ]
            },
            {
                'title': '四、购后行为与产品处置',
                'items': [
                    '购后行为：消费者满意度 → 重复购买 → 口碑传播（正面或负面）。',
                    '产品处置方式：丢弃、赠送、出售、回收。',
                    '物质主义：将物质财富作为幸福和成功的标志。经常在一件旧产品还能使用时购买新产品是物质主义社会的标志之一。',
                ]
            },
        ]
    },

    '第十一章': {
        'title': '群体和社交媒体',
        'key': False,
        'points': [
            {
                'title': '一、参照群体',
                'items': [
                    '参照群体：对消费者行为产生影响的个人或群体。包括：①成员群体——消费者实际归属的群体；②向往群体——消费者希望加入的群体；③疏离群体——消费者想要与之保持距离的群体。',
                    '参照群体的影响方式：信息性影响（提供信息和意见）、规范性影响（遵守群体规范）、价值表达性影响（通过群体认同表达自身价值）。',
                ]
            },
            {
                'title': '二、意见领袖与口碑传播',
                'items': [
                    '意见领袖：具有专业知识或社会影响力，能影响他人购买决策的人。他们往往是早期使用者，在社交媒体上有较多追随者。',
                    '口碑传播（WOM）：消费者之间的非正式信息交流，对购买决策影响巨大。口碑比商业广告更可信。',
                    '网上消费社区：互联网造就的虚拟社区，成员们分享各种内容和购物体验。',
                    '社交媒体营销：利用社交平台（微信、微博、抖音等）建立品牌与消费者的互动。',
                ]
            },
            {
                'title': '三、从众行为与社交网络效应',
                'items': [
                    '从众行为：个体为了符合群体期望而改变自己的行为或态度。',
                    '社交网络效应：消费者的购买决策受其社交网络中其他人的影响。品牌可以通过社交网络扩大影响力。',
                ]
            },
        ]
    },

    '第十二章': {
        'title': '收入和社会阶层',
        'key': False,
        'points': [
            {
                'title': '一、社会阶层',
                'items': [
                    '社会阶层：基于收入、教育、职业等因素对人群的层次划分。同一阶层的人具有相似的价值观、兴趣和行为模式。',
                    '不同阶层的消费差异：上层——注重品味和独特性；中层——追求性价比和品质平衡；下层——注重基本需求满足。',
                    '阶层流动性：社会阶层可以发生变化（向上或向下流动），这影响消费者的品牌偏好和消费模式。',
                ]
            },
            {
                'title': '二、收入与消费',
                'items': [
                    '可支配收入决定消费能力；收入预期影响消费信心。',
                    '炫耀性消费（凡勃伦效应）：为展示社会地位和财富而进行的消费。商品价格越高反而需求越大。',
                    '数字鸿沟：不同社会阶层在获取和使用数字技术方面的差异，影响信息获取和消费方式。',
                ]
            },
        ]
    },

    '第十三章': {
        'title': '亚文化',
        'key': False,
        'points': [
            {
                'title': '一、亚文化的类型',
                'items': [
                    '亚文化定义：在主流文化内具有独特信念、价值观和习俗的子群体。',
                    '民族亚文化：不同民族群体在消费习惯、食物偏好、服饰等方面的差异。',
                    '年龄亚文化：不同年龄段的消费特征差异明显——Z世代（数字原住民）、千禧一代（体验导向）、婴儿潮一代（注重健康和安全）。',
                    '区域亚文化：不同地区消费者的偏好差异，如南北饮食差异、城乡消费差异。',
                ]
            },
            {
                'title': '二、消费体验与体验营销',
                'items': [
                    '消费体验类型：①娱乐体验——以娱乐为主要目的；②教育体验——以学习为主要目的；③审美体验——以美的享受为主要目的；④逃避现实体验——暂时脱离日常生活。',
                    '营造体验的关键要素：①场景设计——与品牌定位契合；②互动参与——提升投入感；③情感共鸣——核心要素；④一致性——各个触点保持统一；⑤个性化——满足个体差异。',
                ]
            },
        ]
    },

    '第十四章': {
        'title': '文化',
        'key': False,
        'points': [
            {
                'title': '一、文化的定义与维度',
                'items': [
                    '文化定义：影响社会成员价值观、信仰、偏好和行为的社会规范系统。文化是习得的，不是天生的。',
                    '文化价值观维度：①个人主义 vs 集体主义；②权力距离（对不平等的接受程度）；③不确定性规避（对模糊和风险的容忍度）；④男性化 vs 女性化（对成就和关怀的重视程度）。',
                    '文化对消费的具体影响：不同文化对颜色、符号、礼仪的解释不同。如白色在西方象征纯洁，在东方部分文化中用于丧礼。',
                ]
            },
            {
                'title': '二、文化变迁与全球化',
                'items': [
                    '文化变迁：社会价值观的变化创造新的消费机会（如健康意识增强带来运动消费增长）。',
                    '全球化与本土化：跨国企业在全球市场面临标准化（降低成本）vs 本土化（适应当地文化）的选择。',
                    '文化适应策略：企业需要适应目标市场的文化特征，包括产品设计、广告内容、渠道选择等。',
                    '文化融合（acculturation）：学习另一种文化所认可的信念与行为的过程。移民、留学、跨国工作等都是文化融合的常见场景。',
                ]
            },
        ]
    },
}

# ============================================================
# STEP 6: Create Word Document
# ============================================================

def set_cell_font(cell, font_name='微软雅黑', font_size=Pt(10)):
    """Set font for all paragraphs in a cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = font_size

def add_formatted_paragraph(doc, text, style=None, font_size=Pt(11), bold=False,
                           color=None, alignment=None, font_name='微软雅黑',
                           space_before=Pt(0), space_after=Pt(4)):
    """Add a paragraph with proper Chinese font handling."""
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()

    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.bold = bold
    if color:
        run.font.color.rgb = color

    pf = p.paragraph_format
    pf.space_before = space_before
    pf.space_after = space_after
    if alignment is not None:
        pf.alignment = alignment

    return p

def add_heading_styled(doc, text, level=1):
    """Add a heading with Microsoft YaHei font."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_bullet_point(doc, text, font_size=Pt(10.5)):
    """Add a bullet point with proper formatting."""
    p = doc.add_paragraph(style='List Bullet')
    # Clear default run and add new one
    p.clear()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = font_size
    return p

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Set page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ============================================================
# COVER PAGE
# ============================================================
# Add empty paragraphs for spacing
for _ in range(6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

add_formatted_paragraph(doc, '14443 消费者行为学',
                       font_size=Pt(28), bold=True,
                       color=RGBColor(0x1a, 0x56, 0x8e),
                       alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       space_after=Pt(8))

add_formatted_paragraph(doc, '《消费者行为学》课程串讲 · 知识梳理',
                       font_size=Pt(16), bold=False,
                       color=RGBColor(0x55, 0x55, 0x55),
                       alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       space_after=Pt(30))

# Separator line
p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━' * 30)
run.font.color.rgb = RGBColor(0x1a, 0x56, 0x8e)
run.font.size = Pt(10)

add_formatted_paragraph(doc, '课程代码：14433 | 适用地区：广东、吉林、山西、福建',
                       font_size=Pt(10.5),
                       color=RGBColor(0x88, 0x88, 0x88),
                       alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       space_before=Pt(20))
add_formatted_paragraph(doc, '基于串讲 PDF 课件全面整理 | 涵盖全部14章核心考点',
                       font_size=Pt(10.5),
                       color=RGBColor(0x88, 0x88, 0x88),
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_formatted_paragraph(doc, '2026年6月',
                       font_size=Pt(10.5),
                       color=RGBColor(0x88, 0x88, 0x88),
                       alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       space_before=Pt(30))

# Page break
doc.add_page_break()

# ============================================================
# EXAM OVERVIEW
# ============================================================
add_heading_styled(doc, '考试概述与应试指导', level=1)

add_formatted_paragraph(doc, '题型分析及答题技巧', font_size=Pt(12), bold=True,
                       color=RGBColor(0x1a, 0x56, 0x8e), space_before=Pt(10), space_after=Pt(6))

exam_tips = [
    '单选题：相对简单，精准选择，注意审题。',
    '多选题：全面思考、宁缺毋滥，不确定的选项不选。',
    '判断题：观点辨析，需能找出错误原因，注意绝对化表述。',
    '名词解释：针对教材中的概念、专业名词进行命题，考核理解能力。答题要简明、概括，突出重点。',
    '简答题：观点突出、明确、逻辑清晰、完整。',
    '论述题：观点突出、明确、逻辑清晰、完整，需有解释说明和例证。',
]
for tip in exam_tips:
    add_bullet_point(doc, tip, font_size=Pt(10.5))

add_formatted_paragraph(doc, '自学方法指导', font_size=Pt(12), bold=True,
                       color=RGBColor(0x1a, 0x56, 0x8e), space_before=Pt(14), space_after=Pt(6))

study_tips = [
    '系统学习，深入重点——先把握整体框架，再深入关键章节。',
    '用心识记，注意领会和理解——不要死记硬背，理解概念背后的逻辑。',
    '运用科学学习方法，注意理论与实践相结合——将知识点与实际消费现象对照。',
    '注意前后的联系和对比——不同章节之间的概念互有关联。',
    '重视实践训练，加强思考——通过做题检验学习效果。',
]
for tip in study_tips:
    add_bullet_point(doc, tip, font_size=Pt(10.5))

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_heading_styled(doc, '目  录', level=1)

toc_items = [
    ('第一章', '购买、拥有和存在：消费者行为学概述', False),
    ('第二章', '消费者与社会福祉', False),
    ('第三章', '知觉', True),
    ('第四章', '学习和记忆', True),
    ('第五章', '动机和情感', True),
    ('第六章', '自我：心智、性别和身体', False),
    ('第七章', '个性、生活方式和价值观', False),
    ('第八章', '态度与劝说沟通', True),
    ('第九章', '制定决策', False),
    ('第十章', '购买、使用与处置', True),
    ('第十一章', '群体和社交媒体', False),
    ('第十二章', '收入和社会阶层', False),
    ('第十三章', '亚文化', False),
    ('第十四章', '文化', False),
]

for ch_num, ch_title, is_key in toc_items:
    text = f'{ch_num}  {ch_title}'
    if is_key:
        text += '  ★ 重点章'
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    if is_key:
        run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)

doc.add_page_break()

# ============================================================
# CHAPTER CONTENT
# ============================================================
chapter_order = ['第一章', '第二章', '第三章', '第四章', '第五章', '第六章', '第七章',
                 '第八章', '第九章', '第十章', '第十一章', '第十二章', '第十三章', '第十四章']

for ch_num in chapter_order:
    ch_data = KNOWLEDGE.get(ch_num)
    if not ch_data:
        continue

    # Chapter heading
    title_text = f'{ch_num}  {ch_data["title"]}'
    if ch_data['key']:
        title_text += '  ★ 重点章'

    h = doc.add_heading(title_text, level=1)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if ch_data['key']:
        for run in h.runs:
            run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)

    for point in ch_data['points']:
        # Sub-heading
        h2 = doc.add_heading(point['title'], level=2)
        for run in h2.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        for item in point['items']:
            add_bullet_point(doc, item, font_size=Pt(10.5))

    # Add page break after each chapter (except last)
    if ch_num != '第十四章':
        doc.add_page_break()

# ============================================================
# FOOTER PAGE
# ============================================================
add_formatted_paragraph(doc, '', font_size=Pt(8), space_before=Pt(20))
p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━' * 30)
run.font.color.rgb = RGBColor(0x1a, 0x56, 0x8e)
run.font.size = Pt(10)

add_formatted_paragraph(doc, '14443 消费者行为学 · 知识梳理 · 基于串讲 PDF 课件整理',
                       font_size=Pt(9), color=RGBColor(0x99, 0x99, 0x99),
                       alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       space_before=Pt(12))
add_formatted_paragraph(doc, '涵盖全部14章核心考点 · 重点章标注 ★',
                       font_size=Pt(9), color=RGBColor(0x99, 0x99, 0x99),
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)

# ============================================================
# SAVE
# ============================================================
output_path = 'D:/ENLIVE/six book/《消费者行为学》知识梳理.docx'
doc.save(output_path)
print(f'\nDocument saved to: {output_path}')
print('Done!')
